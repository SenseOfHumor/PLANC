import pdfplumber
from docx import Document

def read_pdf_file(file_path: str) -> str:
    """
    Reads a PDF file and extracts structured text from all pages using pdfplumber
    for structural accuracy.
    Args:
        file_path (str): The path to the PDF file.
    Returns:
        str: The extracted structured text from the PDF file.
    """
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n\n"
    # print(text)
    return text



def read_docx_file(file_path: str) -> str:
    """
    Reads a DOCX file and extracts text from all paragraphs and tables.
    Args:
        file_path (str): The path to the DOCX file.
    Returns:
        str: The extracted text from the DOCX file.
    """
    doc = Document(file_path)
    text = ""
    # para
    for para in doc.paragraphs:
        text += para.text + "\n"
    # table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\t"
            text = text.rstrip("\t") + "\n"  # Remove the trailing tab and add a newline
    return text

